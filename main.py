import io
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import mysql.connector
from mysql.connector import Error
from pydantic import BaseModel
from passlib.context import CryptContext
from transformers import pipeline, AutoTokenizer, TFMarianMTModel
from typing import List
import torch
import requests
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF
import docx2txt
import chardet
import langid
import google.generativeai as genai
import uvicorn
import mysql
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

 
# app = FastAPI()
app = FastAPI()
host = "cdaserver.mysql.database.azure.com"
user = "cdaadmin"
password = "Qwerty*1"
db = "doctrans"
 
def get_db_connection():
    try:
        connection = mysql.connector.connect(host=host, user=user, password=password, database=db)
        return connection
    except Error as e:
        print("Error while connecting to MySQL", e)
        raise HTTPException(status_code=500, detail="Database connection failed")
 
# Allow CORS for the frontend
origins = [
    "http://127.0.0.1:5500",
    "http://localhost:5500",
    "http://127.0.0.1:5501",
    "https://immune-grouper-marginally.ngrok-free.app",
   
   
    # Add other allowed origins here
]
 
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
 
 
class RegisterRequest(BaseModel):
    username: str
    email: str
    password: str
 
# Password hashing context
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
 
def convertBinarytoFile(binarydata, filename):
    file = io.BytesIO(binarydata)
    file.seek(0)  # Ensure the stream is at the beginning
    return file
@app.get("/", response_class=HTMLResponse)
async def index():
    print('Request for index page received')
    with open("login.html", "r") as file:
        return HTMLResponse(content=file.read(), status_code=200)
@app.get("/download/translated/{translation_id}")
async def download_translated_document(translation_id: int):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
 
        # Query to retrieve the translation based on translation_id
        query = """SELECT t.translated_content, d.document_name, t.language
                   FROM Translations t
                   LEFT JOIN Documents d ON t.original_document_id = d.document_id
                   WHERE t.translation_id = %s"""
        cursor.execute(query, (translation_id,))
        result = cursor.fetchone()
 
        if result:
            translated_content, original_name, language = result
 
             # Convert the binary content to a file-like object
            file_content = convertBinarytoFile(translated_content)
            translated_filename = f"{original_name}_{language[:2].lower()}.pdf"
 
            # Use StreamingResponse to serve the file
            response = StreamingResponse(file_content, media_type='application/pdf')
            response.headers['Content-Disposition'] = f'attachment; filename="{translated_filename}"'
            return response
        else:
            raise HTTPException(status_code=404, detail="Translation not found.")
 
    except mysql.connector.Error as e:
        print("Failed to retrieve translation.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve translation.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
       
# Registration endpoint
@app.post("/register/")
async def register(request: RegisterRequest):
    connection = get_db_connection()
    cursor = connection.cursor()
 
    # Hash the user's password
    hashed_password = pwd_context.hash(request.password)
 
    try:
        insert_user_query = """
        INSERT INTO Users (username, password, email)
        VALUES (%s, %s, %s)
        """
        cursor.execute(insert_user_query, (request.username, hashed_password, request.email))
        connection.commit()
        return {"message": "User registered successfully"}
 
    except Error as e:
        connection.rollback()
        if e.errno == 1062:  # Duplicate entry error
            raise HTTPException(status_code=400, detail="Username already exists.")
        else:
            raise HTTPException(status_code=500, detail="Internal server error.")
   
    finally:
        cursor.close()
        connection.close()
       
# Function to store document in the database
def store_document(user_id, document_name, document_content):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
       
        insert_doc_query = """
        INSERT INTO Documents (user_id, document_name, document_content)
        VALUES (%s, %s, %s)
        """
        cursor.execute(insert_doc_query, (user_id, document_name, document_content))
        connection.commit()
        print(cursor.lastrowid)
        return cursor.lastrowid
   
    except Error as e:
        print("Failed to store document.", e)
        raise HTTPException(status_code=500, detail="Failed to store document.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
 
 
 
def store_translation(original_document_id, language, translated_content):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
 
        # First, execute the SELECT query
        select_document_query = "SELECT * FROM Documents WHERE document_id = %s"
        cursor.execute(select_document_query, (original_document_id,))
       
        # Fetch all the results (even if you don't need them)
        cursor.fetchall()
 
        # Now insert the translation into the Translations table
        insert_translation_query = """
        INSERT INTO Translations (original_document_id, language, translated_content)
        VALUES (%s, %s, %s)
        """
        cursor.execute(insert_translation_query, (original_document_id, language, translated_content))
        connection.commit()
       
        return cursor.lastrowid
 
    except mysql.connector.Error as e:
        print("Failed to store translation.", e)
        raise HTTPException(status_code=500, detail="Failed to store translation.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
 
 
 
class ForgotPasswordRequest(BaseModel):
    email: str
 
 
class LoginRequest(BaseModel):
    username: str
    password: str
 
 
 
 
@app.post("/forgot-password/")
async def forgot_password(request: ForgotPasswordRequest):
    connection = get_db_connection()
    cursor = connection.cursor(buffered=True)  # Use a buffered cursor
 
    try:
        # Check if the email exists in the Users table
        query = "SELECT user_id FROM Users WHERE email = %s"
        cursor.execute(query, (request.email,))
        result = cursor.fetchone()
 
        if result is None:
            raise HTTPException(status_code=404, detail="Email not found")
 
        # Implement your password reset logic here, e.g., generate a reset token and send an email
        # For now, we will just return a success message
        return {"message": "Password reset instructions have been sent to your email."}
 
    except Error as e:
        raise HTTPException(status_code=500, detail="Failed to process request.")
   
    finally:
        cursor.close()
        connection.close()
 
# Login endpoint
@app.post("/login/")
async def login(request: LoginRequest):
    connection = get_db_connection()
    cursor = connection.cursor()
 
    try:
        query = "SELECT password FROM Users WHERE username = %s"
        cursor.execute(query, (request.username,))
        result = cursor.fetchone()
        print(result)
 
        if result is None or not pwd_context.verify(request.password, result[0]):
            raise HTTPException(status_code=401, detail="Invalid username or password")
 
        return {"message": "Login successful"}
 
    finally:
        cursor.close()
        connection.close()
 
 
def convertFileToBinary(file):
    binarydata = file.read()
    return binarydata
 
# Endpoint to handle document upload and translation
 
@app.post("/upload-original-document/")
async def uploadoriginaldocument(user_id: int = Form(...), language: str = Form(...), file: UploadFile = File(...)):
    # file_content = await file.read()
    binary_content = convertFileToBinary(file.file)
    document_id = store_document(user_id, file.filename, binary_content)
 
    return JSONResponse(content={"document_id": document_id})
 
 
 
 
@app.post("/upload-translated-document/")
async def uploadtranslateddocument(doc_id: int = Form(...), language: str = Form(...), file: UploadFile = File(...)):
    # file_content = await file.read()
    binary_content = convertFileToBinary(file.filename)
    document_id = store_translation(doc_id, file.filename, binary_content)
 
    return JSONResponse(content={"document_id": document_id})
# import base64
@app.get("/alldocuments/")
def get_documents_and_translations():
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
 
        # Query to retrieve document details along with their starred status
        select_documents_query = """
        SELECT d.document_name, d.uploaded_at, d.starred, t.language, t.translated_content, t.original_document_id, t.translation_id
        FROM Translations t
        LEFT JOIN Documents d ON t.original_document_id = d.document_id
        """
        cursor.execute(select_documents_query)
        translations = cursor.fetchall()
 
        if translations:
            document_dict = {}
            for translation in translations:
                doc_name, uploaded_at, starred, lang, trans_content, original_id, translated_id = translation
                if doc_name not in document_dict:
                    document_dict[doc_name] = {
                        "name": doc_name,
                        "translated": [],
                        "sourceLang": "English",
                        "translatedLang": [],
                        "uploadedDate": uploaded_at.strftime("%Y-%m-%d"),
                        "starred": bool(starred),  # Get the starred status from the DB
                        "docID": original_id,
                        "transID": translated_id
                    }
                if lang:
                    translated_doc_name = f"{lang}"
                    document_dict[doc_name]["translated"].append(translated_doc_name)
                    language = lang.split('_')[1].split('.')[0]
                    document_dict[doc_name]["translatedLang"].append(language)
 
            # Format the data to match the frontend structure
            documents_list = []
            for doc_name, details in document_dict.items():
                documents_list.append({
                    "name": details["name"],
                    "translated": ", ".join(details["translated"]),
                    "sourceLang": details["sourceLang"],
                    "translatedLang": ", ".join(details["translatedLang"]),
                    "uploadedDate": details["uploadedDate"],
                    "starred": details["starred"],
                    "docID": details["docID"],
                    "transID": details["transID"]
                })
            return documents_list
        else:
            raise HTTPException(status_code=404, detail="No translations found.")
 
    except Error as e:
        print("Failed to retrieve translations.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve translations.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
              
@app.get("/stareddocuments/")
def get_starred_documents_and_translations():
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
 
        # Query to retrieve starred translations and their corresponding original document names and upload dates
        select_documents_query = """
        SELECT d.document_name, d.uploaded_at, t.language, t.translated_content, t.original_document_id, t.translation_id, d.starred
        FROM Translations t
        LEFT JOIN Documents d ON t.original_document_id = d.document_id
        WHERE d.starred = TRUE
        """
        cursor.execute(select_documents_query)
        translations = cursor.fetchall()
 
        if translations:
            document_dict = {}
            for translation in translations:
                doc_name, uploaded_at, lang, trans_content, original_id, translated_id, starred = translation
                if doc_name not in document_dict:
                    document_dict[doc_name] = {
                        "name": doc_name,
                        "translated": [],
                        "sourceLang": "English",
                        "translatedLang": [],
                        "uploadedDate": uploaded_at.strftime("%Y-%m-%d"),
                        "starred": starred,
                        "docID": original_id,
                        "transID": translated_id
                    }
                if lang:
                    translated_doc_name = f"{lang}"
                    document_dict[doc_name]["translated"].append(translated_doc_name)
                    filename = lang
                    language = filename.split('_')[1].split('.')[0]
                    document_dict[doc_name]["translatedLang"].append(language)
 
            # Format the data to match the frontend structure
            documents_list = []
            for doc_name, details in document_dict.items():
                documents_list.append({
                    "name": details["name"],
                    "translated": ", ".join(details["translated"]),
                    "sourceLang": details["sourceLang"],
                    "translatedLang": ", ".join(details["translatedLang"]),
                    "uploadedDate": details["uploadedDate"],
                    "starred": details["starred"],
                    "docID": details["docID"],
                    "transID": details["transID"]
                })
 
            return documents_list
        else:
            raise HTTPException(status_code=404, detail="No starred documents found.")
 
    except Error as e:
        print("Failed to retrieve translations.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve translations.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
 
# Function to download original document
@app.get("/download/original/{document_id}")
async def download_original_document(document_id: int):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
 
        # Query to retrieve the document based on document_id
        query = "SELECT document_name, document_content FROM Documents WHERE document_id = %s"
        cursor.execute(query, (document_id,))
        result = cursor.fetchone()
 
        if result:
            document_name, document_content = result
 
            # Convert binary content to bytes and prepare it for download
            file_content = io.BytesIO(document_content)
 
            return FileResponse(file_content, media_type='application/octet-stream', filename=document_name)
        else:
            raise HTTPException(status_code=404, detail="Document not found.")
 
    except Error as e:
        print("Failed to retrieve document.", e)
        raise HTTPException(status_code=500, detail="Failed to retrieve document.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()
           
 
 
@app.post("/documents/toggle_star/{document_id}")
async def toggle_star(document_id: int):
    try:
        connection = get_db_connection()
        cursor = connection.cursor()
 
        # Retrieve the current starred status
        query = "SELECT starred FROM Documents WHERE document_id = %s"
        cursor.execute(query, (document_id,))
        result = cursor.fetchone()
 
        if result is None:
            raise HTTPException(status_code=404, detail="Document not found.")
 
        # Toggle the starred status
        current_status = result[0]
        new_status = not current_status
 
        # Update the document's starred status in the database
        update_query = "UPDATE Documents SET starred = %s WHERE document_id = %s"
        cursor.execute(update_query, (new_status, document_id))
        connection.commit()
 
        return {"starred": new_status}
 
    except Error as e:
        print("Failed to toggle star status.", e)
        raise HTTPException(status_code=500, detail="Failed to toggle star status.")
   
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()



class TextPayload(BaseModel):
    text: str
 
# Mapping of language codes to full names
language_map = {
    "af": "Afrikaans",
    "am": "Amharic",
    "an": "Aragonese",
    "ar": "Arabic",
    "as": "Assamese",
    "az": "Azerbaijani",
    "be": "Belarusian",
    "bg": "Bulgarian",
    "bn": "Bengali",
    "br": "Breton",
    "bs": "Bosnian",
    "ca": "Catalan",
    "cs": "Czech",
    "cy": "Welsh",
    "da": "Danish",
    "de": "German",
    "dz": "Dzongkha",
    "el": "Greek",
    "en": "English",
    "eo": "Esperanto",
    "es": "Spanish",
    "et": "Estonian",
    "eu": "Basque",
    "fa": "Persian",
    "fi": "Finnish",
    "fo": "Faroese",
    "fr": "French",
    "ga": "Irish",
    "gl": "Galician",
    "gu": "Gujarati",
    "he": "Hebrew",
    "hi": "Hindi",
    "hr": "Croatian",
    "ht": "Haitian Creole",
    "hu": "Hungarian",
    "hy": "Armenian",
    "id": "Indonesian",
    "is": "Icelandic",
    "it": "Italian",
    "ja": "Japanese",
    "jv": "Javanese",
    "ka": "Georgian",
    "kk": "Kazakh",
    "km": "Khmer",
    "kn": "Kannada",
    "ko": "Korean",
    "ku": "Kurdish",
    "ky": "Kyrgyz",
    "la": "Latin",
    "lb": "Luxembourgish",
    "lo": "Lao",
    "lt": "Lithuanian",
    "lv": "Latvian",
    "mg": "Malagasy",
    "mk": "Macedonian",
    "ml": "Malayalam",
    "mn": "Mongolian",
    "mr": "Marathi",
    "ms": "Malay",
    "mt": "Maltese",
    "nb": "Norwegian Bokmål",
    "ne": "Nepali",
    "nl": "Dutch",
    "nn": "Norwegian Nynorsk",
    "no": "Norwegian",
    "oc": "Occitan",
    "or": "Oriya",
    "pa": "Punjabi",
    "pl": "Polish",
    "ps": "Pashto",
    "pt": "Portuguese",
    "qu": "Quechua",
    "ro": "Romanian",
    "ru": "Russian",
    "rw": "Kinyarwanda",
    "se": "Northern Sami",
    "si": "Sinhala",
    "sk": "Slovak",
    "sl": "Slovenian",
    "sq": "Albanian",
    "sr": "Serbian",
    "sv": "Swedish",
    "sw": "Swahili",
    "ta": "Tamil",
    "te": "Telugu",
    "th": "Thai",
    "tl": "Tagalog",
    "tr": "Turkish",
    "ug": "Uighur",
    "uk": "Ukrainian",
    "ur": "Urdu",
    "vi": "Vietnamese",
    "vo": "Volapük",
    "wa": "Walloon",
    "xh": "Xhosa",
    "zh": "Chinese",
    "zu": "Zulu"
}

def extract_text_from_file(file: UploadFile) -> str:
    try:
        if file.content_type == "text/plain":
            content = file.file.read()
            result = chardet.detect(content)
            text = content.decode(result['encoding'])
        elif file.content_type == "application/pdf":
            pdf_reader = PdfReader(BytesIO(file.file.read()))
            text = "".join([page.extract_text() for page in pdf_reader.pages])
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = docx2txt.process(BytesIO(file.file.read()))
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")

@app.post("/detect-language/")
async def detect_language(file: UploadFile = File(...)):
    try:
        text = extract_text_from_file(file)
        if not text:
            raise HTTPException(status_code=400, detail="No text extracted from the file")

        language_code, confidence = langid.classify(text)
        if language_code in language_map:
            return {"language": language_map[language_code]}
        else:
            raise HTTPException(status_code=400, detail="No language detected")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 

# Configure the Google AI SDK
genai.configure(api_key="AIzaSyC2ykmE8kwmam1D6Xn60l7XE964CuTKGS8")

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config
)

@app.post("/upload-summarize-translate/")
async def upload_summarize_translate(language: str = Form(...), 
                                     file: UploadFile = File(...), 
                                     output_format: str = Form(...)):
    try:
        print("language received: ", language)
        print("File received: ", file)
        print("Output format: ", output_format)

        # Determine file type and extract text
        file_extension = file.filename.split(".")[-1].lower()
        text = ""

        if file_extension == "pdf":
            reader = PdfReader(file.file)
            for page in reader.pages:
                text += page.extract_text()
        elif file_extension == "txt":
            text = (await file.read()).decode("utf-8")
        elif file_extension == "docx":
            doc = Document(file.file)
            for para in doc.paragraphs:
                text += para.text
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")

        # Limit the extracted text to 400 words
        words = text.split()
        limited_text = " ".join(words[:400])

        # Prepare the prompt for summarization
        prompt = f"summarize the following text up to 400 words in {language}: {limited_text}"

        # Start a chat session with the Gemini LLM
        chat_session = model.start_chat(history=[
            {
                "role": "user",
                "parts": [prompt],
            }
        ])

        # Get the summary from the model
        summary_response = chat_session.send_message(prompt)
        summary = summary_response.text

        # Remove newline characters from the summary
        summary = summary.replace('\n', ' ')

        # Save the summary in the desired file format
        output_file_path = ""
        if output_format.lower() == "txt":
            output_file_path = "summary.txt"
            with open(output_file_path, "w", encoding="utf-8") as f:
                f.write(summary)
        elif output_format.lower() == "docx":
            output_file_path = "summary.docx"
            doc = Document()
            doc.add_paragraph(summary)
            doc.save(output_file_path)
        elif output_format.lower() == "pdf":
            output_file_path = "summary.pdf"
            
            # Use fpdf2 and set the path to the Unicode font
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font("DejaVu", '', './fonts/DejaVuSansCondensed.ttf', uni=True)  # Provide the correct path
            pdf.set_font("DejaVu", size=12)
            pdf.multi_cell(0, 10, summary)
            pdf.output(output_file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported output format")

        # Return the saved file as a binary stream
        def iterfile():
            with open(output_file_path, "rb") as f:
                yield from f

        return StreamingResponse(iterfile(), 
                                 media_type="application/octet-stream", 
                                 headers={
                                     "Content-Disposition": f"attachment; filename={output_file_path}"
                                 })

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

if __name__ == "__main__":
    uvicorn.run("backend:app", host="0.0.0.0", port=8000)
